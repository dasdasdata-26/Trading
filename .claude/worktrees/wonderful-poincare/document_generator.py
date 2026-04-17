"""
HPM Systems Proposal Generator — .docx generation module.
Produces output visually matching the sample PDF (226xxx-Client_Title_rev#.pdf).
"""

import io
import re
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI

from bs4 import BeautifulSoup, NavigableString, Tag

BASE_DIR = Path(__file__).parent
LOGO_PATH = BASE_DIR / "static" / "img" / "HPM logo.png"

ROMAN = ["I", "II", "III", "IV"]
LETTERS = list("ABCDEFGHIJKLMNOPQRSTUVWXYZ")

# Twips (1 inch = 1440 twips). Text width = 6.5" = 9360 twips
TEXT_WIDTH_TWIPS = 9360


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def generate_proposal(data, settings):
    """Generate a .docx proposal and return bytes."""
    doc = Document()

    # Remove the default blank paragraph Word always adds
    for p in doc.paragraphs:
        _remove_para(p)

    # Apply global Normal style defaults
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    # Fix 1: enforce single spacing on Normal style so all paragraphs inherit it
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1.4)
    section.bottom_margin = Inches(1)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)
    section.different_first_page_header_footer = True

    # Headers and footer
    _build_first_page_header(section)
    _build_continuation_header(section, data)
    _build_footer(section, settings)
    _build_first_page_footer(section, settings)

    # Body
    _build_body(doc, data, settings)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Headers / Footer
# ---------------------------------------------------------------------------

def _build_first_page_header(section):
    """Page 1 header: HPM logo, left-aligned."""
    hdr = section.first_page_header
    # Clear default content
    for p in hdr.paragraphs:
        _remove_para(p)

    p = hdr.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    if LOGO_PATH.exists():
        run.add_picture(str(LOGO_PATH), width=Inches(1.5))
    else:
        run.text = "HPM Systems, Inc."
        _fmt(run, 12, bold=True)


def _build_continuation_header(section, data):
    """Pages 2+ header: 7pt Arial continuation block."""
    hdr = section.header
    for p in hdr.paragraphs:
        _remove_para(p)

    date_str = _format_date(data.get("date", ""))
    lines = [
        data.get("company_name", ""),
        f"Subject: {data.get('subject_line', '')}",
        f"Site: {data.get('site_location', '')}",
        f"Proposal {data.get('proposal_number', '')}",
        date_str,
    ]

    for line in lines:
        p = hdr.add_paragraph()
        run = p.add_run(line)
        _fmt(run, 7)
        _para_spacing(p, 0, 0)

    # "Page X of Y" line
    p = hdr.add_paragraph()
    _para_spacing(p, 0, 0)
    r = p.add_run("Page ")
    _fmt(r, 7)
    _add_field_run(p, "PAGE", size=7)
    r = p.add_run(" of ")
    _fmt(r, 7)
    _add_field_run(p, "NUMPAGES", size=7)


def _build_footer(section, settings):
    """All-pages footer: centered 7pt address line."""
    ftr = section.footer
    for p in ftr.paragraphs:
        _remove_para(p)
    footer_text = settings.get("footer", {}).get(
        "text", "1292 Kifer Road, Suite 807, Sunnyvale, CA 94086  \u2666  Ph: (408) 615-6900"
    )
    p = ftr.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(footer_text)
    _fmt(run, 7)
    _para_spacing(p, 0, 0)


def _build_first_page_footer(section, settings):
    """First-page footer: same as default footer."""
    ftr = section.first_page_footer
    for p in ftr.paragraphs:
        _remove_para(p)
    footer_text = settings.get("footer", {}).get(
        "text", "1292 Kifer Road, Suite 807, Sunnyvale, CA 94086  \u2666  Ph: (408) 615-6900"
    )
    p = ftr.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(footer_text)
    _fmt(run, 7)
    _para_spacing(p, 0, 0)


# ---------------------------------------------------------------------------
# Body
# ---------------------------------------------------------------------------

def _build_body(doc, data, settings):
    date_str = _format_date(data.get("date", ""))
    proposal_num = data.get("proposal_number", "")
    first_name = data.get("recipient_first_name", "")
    last_name = data.get("recipient_last_name", "")
    company = data.get("company_name", "")
    street = data.get("street_address", "")
    city_state_zip = data.get("city_state_zip", "")
    subject = data.get("subject_line", "")
    site = data.get("site_location", "")

    user = settings.get("user", {})
    company_info = settings.get("company", {})

    # --- Date + Proposal number ---
    p = doc.add_paragraph()
    _para_spacing(p, 0, 3)
    _add_right_tab(p, TEXT_WIDTH_TWIPS)
    r1 = p.add_run(date_str)
    _fmt(r1, 9, bold=True)
    r2 = p.add_run(f"\tProposal: {proposal_num}")
    _fmt(r2, 9)

    # --- Addressee block ---
    _blank(doc)
    _body_para(doc, f"{first_name} {last_name}".strip())
    _body_para(doc, company)
    _body_para(doc, street)
    _body_para(doc, city_state_zip)

    # --- RE: block ---
    _blank(doc)
    _body_para(doc, f"RE: Subject: {subject}")
    _body_para(doc, f"Site: {site}")

    # --- Salutation ---
    _blank(doc)
    _body_para(doc, f"Dear {first_name},")

    # --- Intro paragraph ---
    _blank(doc)
    intro = data.get("intro_paragraph", "")
    _body_para(doc, intro)

    # Set up Word list numbering (Fix 4)
    num_state = _setup_numbering(doc)

    # --- SCOPE OF WORK heading ---
    _blank(doc)
    p = doc.add_paragraph()
    _para_spacing(p, 0, 0)
    r = p.add_run("SCOPE OF WORK")
    _fmt(r, 9, bold=True, underline=True)

    # --- Sections I–IV (Fix 2, 3, 4) ---
    sections = data.get("sections", [])
    # Fixes 2 & 3: skip deleted sections, renumber remaining ones sequentially
    active_sections = [(orig_idx, sec) for orig_idx, sec in enumerate(sections)
                       if not sec.get("deleted")]

    for visual_idx, (orig_idx, sec) in enumerate(active_sections):
        heading = sec.get("heading", "")
        items = sec.get("items", [])

        # Visual blank before each section heading (replaces space_before=4pt)
        _blank(doc)

        # Section heading: numeral at 0", heading text at 0.5"
        # Build pPr XML in correct OOXML schema order so Word honours the tab stop.
        # Order required: tabs → spacing → ind
        p = doc.add_paragraph()
        _para_spacing(p, 0, 0)   # adds <w:spacing> to pPr
        pPr = p._p.get_or_add_pPr()

        # Insert <w:tabs> BEFORE the existing <w:spacing> element
        tabs_el = OxmlElement("w:tabs")
        tab_el = OxmlElement("w:tab")
        tab_el.set(qn("w:val"), "left")
        tab_el.set(qn("w:pos"), "720")   # 0.5" = 720 twips
        tabs_el.append(tab_el)
        spacing_el = pPr.find(qn("w:spacing"))
        if spacing_el is not None:
            spacing_el.addprevious(tabs_el)
        else:
            pPr.append(tabs_el)

        # Append <w:ind> AFTER <w:spacing> (correct schema order)
        ind_el = OxmlElement("w:ind")
        ind_el.set(qn("w:left"), "720")    # continuation text at 0.5"
        ind_el.set(qn("w:hanging"), "720") # first line (numeral) at 0"
        pPr.append(ind_el)

        r = p.add_run(f"{ROMAN[visual_idx]})\t{heading}")
        _fmt(r, 9, bold=True)

        # Allocate a fresh letter numId for this section so A resets (Fix 4)
        letter_num_id = _alloc_letter_num(num_state)

        for item in items:
            html = item.get("html", "")
            subitems = item.get("subitems", [])

            # A at 0.5", text at 1.0" (hanging=0.5")
            _add_list_para(doc, html, letter_num_id,
                           left_inches=1.0, hanging_inches=0.5)

            if subitems:
                # Fresh number numId for each item's sub-list so 1 resets (Fix 4)
                number_num_id = _alloc_number_num(num_state)
                for subitem in subitems:
                    si_html = subitem.get("html", "")
                    # 1 at 1.0", text at 1.5" (hanging=0.5")
                    _add_list_para(doc, si_html, number_num_id,
                                   left_inches=1.5, hanging_inches=0.5)

    # --- ASSUMPTIONS/EXCLUSIONS ---
    _blank(doc)
    p = doc.add_paragraph()
    _para_spacing(p, 0, 0)
    r = p.add_run("ASSUMPTIONS/EXCLUSIONS")
    _fmt(r, 9, bold=True, underline=True)

    _body_para(doc, "The scope of work is based on the following assumptions and/or exclusions:")

    assumptions = data.get("assumptions", [])
    if assumptions:
        # Fresh letter numId for assumptions list (Fix 4)
        assump_num_id = _alloc_letter_num(num_state)
        for assumption in assumptions:
            _add_list_para(doc, assumption, assump_num_id,
                           left_inches=1.0, hanging_inches=0.5)

    # --- PROFESSIONAL FEE ---
    _blank(doc)
    p = doc.add_paragraph()
    _para_spacing(p, 0, 0)
    r = p.add_run("PROFESSIONAL FEE")
    _fmt(r, 9, bold=True, underline=True)

    fees = data.get("fees", {})
    fee_values = [
        fees.get("section1", 0),
        fees.get("section2", 0),
        fees.get("section3", 0),
        fees.get("section4", 0),
    ]
    # Fix 3: only sum fees for active (non-deleted) sections
    total = sum(_parse_fee(fee_values[orig_idx])
                for orig_idx, _ in active_sections
                if orig_idx < len(fee_values))

    for visual_idx, (orig_idx, sec) in enumerate(active_sections):
        heading = sec.get("heading", "")
        amount = _parse_fee(fee_values[orig_idx]) if orig_idx < len(fee_values) else 0
        label = f"{ROMAN[visual_idx]}) {heading}"
        _add_fee_row(doc, label, amount, bold=False, yellow=False)

    _add_fee_row(doc, "TOTAL FEE", total, bold=True, yellow=True)

    # --- Closing paragraph ---
    _blank(doc)
    closing = data.get("closing_paragraph", "")
    _body_para(doc, closing)

    # --- Signature block ---
    _blank(doc)
    _build_signature(doc, data, settings)

    # --- Rate schedule (new page) ---
    _add_page_break(doc)
    _build_rate_schedule(doc, settings)


# ---------------------------------------------------------------------------
# Signature block (2-column borderless table)
# ---------------------------------------------------------------------------

def _build_signature(doc, data, settings):
    user = settings.get("user", {})
    company_info = settings.get("company", {})
    include_cc = data.get("include_cc", True)
    cc_line = data.get("cc_line", user.get("default_cc", ""))

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    _remove_table_borders(tbl)

    left = tbl.cell(0, 0)
    right = tbl.cell(0, 1)

    # Set column widths
    left.width = Inches(3.5)
    right.width = Inches(3.0)

    # --- Left column: build procedurally to insert signature image ---

    # "Sincerely,"
    p = left.add_paragraph()
    _para_spacing(p, 0, 1)
    r = p.add_run("Sincerely,")
    _fmt(r, 9, bold=False)

    # Signature image (between "Sincerely," and name)
    sig_path_str = user.get("signature_image", "")
    sig_path = None
    if sig_path_str:
        candidate = BASE_DIR / sig_path_str
        if candidate.exists():
            sig_path = candidate
        elif Path(sig_path_str).exists():
            sig_path = Path(sig_path_str)

    if sig_path:
        p = left.add_paragraph()
        _para_spacing(p, 0, 1)
        r = p.add_run()
        r.add_picture(str(sig_path), width=Inches(1.0), height=Inches(0.4))
    else:
        # Fallback: 3 blank lines for manual signature space
        for _ in range(3):
            p = left.add_paragraph()
            _para_spacing(p, 0, 1)
            r = p.add_run("")
            _fmt(r, 9, bold=False)

    # Remaining left-column lines
    remaining_left = [
        (user.get("name", ""), False),
        (user.get("title", ""), False),
        ("HPM Systems, Inc.", False),
        (f"CA contractor license # {company_info.get('contractor_license', '')}", False),
    ]
    if include_cc and cc_line:
        remaining_left.append((cc_line, False))

    for text, bold in remaining_left:
        p = left.add_paragraph()
        _para_spacing(p, 0, 1)
        r = p.add_run(text)
        _fmt(r, 9, bold=bold)

    # --- Right column ---
    right_lines = [
        ("Client", False),
        ("Authorized Representative", False),
        ("", False),
        ("", False),
        ("Signature                                    Date", False),
        ("", False),
        ("P.O.# (if applicable)", False),
    ]

    for text, bold in right_lines:
        p = right.add_paragraph()
        _para_spacing(p, 0, 1)
        r = p.add_run(text)
        _fmt(r, 9, bold=bold)

    # Remove the default empty paragraph that python-docx adds to each cell
    for cell in [left, right]:
        if cell.paragraphs[0].text == "" and len(cell.paragraphs) > 1:
            _remove_para(cell.paragraphs[0])


# ---------------------------------------------------------------------------
# Rate Schedule
# ---------------------------------------------------------------------------

def _build_rate_schedule(doc, settings):
    rs = settings.get("rate_schedule", {})
    year = rs.get("year", "2026")

    # Title
    p = doc.add_paragraph()
    _para_spacing(p, 0, 4)
    r = p.add_run("HPM SYSTEMS RATE SCHEDULE")
    _fmt(r, 10, bold=True)

    _body_para(doc, f"The following fee schedule applies through December 31, {year}")
    _blank(doc)

    # Professional Fees
    p = doc.add_paragraph()
    _para_spacing(p, 4, 2)
    r = p.add_run("PROFESSIONAL FEES")
    _fmt(r, 9, bold=True)

    for fee in rs.get("professional_fees", []):
        role = fee.get("role", "")
        rate = fee.get("rate")
        note = fee.get("note", "")
        if rate is not None:
            text = f"{role}: ${rate:,}"
        else:
            text = f"{role} {note}"
        _body_para(doc, text)

    # Special Charges
    _blank(doc)
    p = doc.add_paragraph()
    _para_spacing(p, 4, 2)
    r = p.add_run("SPECIAL CHARGES")
    _fmt(r, 9, bold=True)

    for charge in rs.get("special_charges", []):
        p = doc.add_paragraph()
        _para_spacing(p, 0, 0)
        pf = p.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.first_line_indent = Inches(-0.25)
        rb = p.add_run("\u2022  ")
        _fmt(rb, 9)
        r = p.add_run(charge)
        _fmt(r, 9)

    # Other Charges (Calibration Gas, Service Van)
    _blank(doc)
    for item in rs.get("other_charges", []):
        _body_para(doc, f"{item['label']}:  {item['charge']}")

    # Reimbursable Expenses
    _blank(doc)
    reimb = rs.get("reimbursable", {})
    p = doc.add_paragraph()
    _para_spacing(p, 4, 2)
    r = p.add_run("REIMBURSABLE EXPENSES")
    _fmt(r, 9, bold=True)

    markup = reimb.get("markup", "")
    intro = reimb.get("intro", "")
    if markup:
        _body_para(doc, f"Direct Expenses {markup}")
    if intro:
        _body_para(doc, intro)

    _body_para(doc, "Plotting and Reproductions (In-House)")
    for item in reimb.get("items", []):
        _body_para(doc, f"    {item['label']}:  {item['charge']}")


# ---------------------------------------------------------------------------
# Helpers: Paragraphs
# ---------------------------------------------------------------------------

def _body_para(doc, text, bold=False):
    """Add a normal 9pt Arial paragraph."""
    p = doc.add_paragraph()
    _para_spacing(p, 0, 2)
    r = p.add_run(text)
    _fmt(r, 9, bold=bold)
    return p


def _blank(doc):
    """Add an empty spacing paragraph."""
    p = doc.add_paragraph()
    _para_spacing(p, 0, 0)
    r = p.add_run("")
    _fmt(r, 4)  # tiny font keeps the line height small
    return p


def _add_html_paragraph(doc, label, html_or_text, indent_left=0.5, indent_hanging=0.3):
    """
    Add a paragraph with label prefix and body content from HTML.
    Handles inline bold/italic/underline and <ul> bullet sub-lists.
    """
    soup = BeautifulSoup(html_or_text, "html.parser")

    # Check for top-level <ul> to decide if we need multiple paragraphs
    # We'll process text nodes + inline formatting first, then bullets
    p = doc.add_paragraph()
    _para_spacing(p, 0, 2)
    pf = p.paragraph_format
    pf.left_indent = Inches(indent_left)
    pf.first_line_indent = Inches(-indent_hanging)

    # Label run
    r = p.add_run(label)
    _fmt(r, 9)

    # Process soup children
    _process_nodes(p, soup.children)

    # Now handle any <ul> elements — add as separate bullet paragraphs
    for ul in soup.find_all("ul"):
        for li in ul.find_all("li"):
            bp = doc.add_paragraph()
            _para_spacing(bp, 0, 1)
            bpf = bp.paragraph_format
            bpf.left_indent = Inches(indent_left + 0.25)
            bpf.first_line_indent = Inches(-0.2)
            br = bp.add_run("\u2022\u2009")
            _fmt(br, 9)
            _process_nodes(bp, li.children)

    return p


def _process_nodes(para, nodes, bold=False, italic=False, underline=False):
    """Recursively process BeautifulSoup nodes into paragraph runs."""
    for node in nodes:
        if isinstance(node, NavigableString):
            text = str(node)
            # Skip if this text belongs to a <ul> (will be handled separately)
            if node.parent and node.parent.name in ("ul", "li"):
                continue
            if text:
                r = para.add_run(text)
                _fmt(r, 9, bold=bold, italic=italic, underline=underline)
        elif isinstance(node, Tag):
            if node.name in ("ul", "li"):
                # ul/li handled by _add_html_paragraph caller
                continue
            nb = bold or node.name in ("strong", "b")
            ni = italic or node.name in ("em", "i")
            nu = underline or node.name == "u"
            if node.name == "br":
                # line break within a paragraph — just a space
                r = para.add_run(" ")
                _fmt(r, 9)
            else:
                _process_nodes(para, node.children, nb, ni, nu)


def _add_fee_row(doc, label, amount, bold=False, yellow=False):
    """Add a fee row paragraph with dot-leader tab stop."""
    p = doc.add_paragraph()
    _para_spacing(p, 0, 2)

    if yellow:
        _set_para_shading(p, "FFFF00")

    # Dot-leader tab at right margin
    _add_dot_leader_tab(p, TEXT_WIDTH_TWIPS)

    r1 = p.add_run(label)
    _fmt(r1, 9, bold=bold)

    r2 = p.add_run(f"\t${amount:,}")
    _fmt(r2, 9, bold=bold)


def _add_page_break(doc):
    """Insert a page break paragraph."""
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    _para_spacing(p, 0, 0)
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


# ---------------------------------------------------------------------------
# Helpers: XML / Formatting
# ---------------------------------------------------------------------------

def _fmt(run, size, bold=False, italic=False, underline=False):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline


def _para_spacing(para, before_pt=0, after_pt=0):
    """Fix 1: single line spacing, 0 pt space before/after on every paragraph.
    Uses direct XML to guarantee w:lineRule='auto' and w:line=240 (single spacing).
    """
    pPr = para._p.get_or_add_pPr()
    existing = pPr.find(qn("w:spacing"))
    if existing is not None:
        pPr.remove(existing)
    spacing_el = OxmlElement("w:spacing")
    spacing_el.set(qn("w:before"), "0")
    spacing_el.set(qn("w:after"), "0")
    spacing_el.set(qn("w:line"), "240")
    spacing_el.set(qn("w:lineRule"), "auto")
    pPr.append(spacing_el)


def _add_right_tab(para, pos_twips):
    """Add a right-aligned tab stop at pos_twips (no leader)."""
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab)
    pPr.append(tabs)


def _add_left_tab(para, pos_twips):
    """Add a left-aligned tab stop at pos_twips."""
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab)
    pPr.append(tabs)


def _add_dot_leader_tab(para, pos_twips):
    """Add a right-aligned tab stop with dot leader at pos_twips."""
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab)
    pPr.append(tabs)


def _set_para_shading(para, fill_hex):
    """Apply background shading to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _add_field_run(para, field_code, size=9):
    """Insert a Word field (PAGE, NUMPAGES, etc.) into a paragraph."""
    # begin
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    rPr.append(szCs)
    r.append(rPr)
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "begin")
    r.append(fc)
    para._p.append(r)

    # instrText
    r = OxmlElement("w:r")
    rPr2 = OxmlElement("w:rPr")
    sz2 = OxmlElement("w:sz")
    sz2.set(qn("w:val"), str(int(size * 2)))
    szCs2 = OxmlElement("w:szCs")
    szCs2.set(qn("w:val"), str(int(size * 2)))
    rPr2.append(sz2)
    rPr2.append(szCs2)
    r.append(rPr2)
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = f" {field_code} "
    r.append(it)
    para._p.append(r)

    # end
    r = OxmlElement("w:r")
    rPr3 = OxmlElement("w:rPr")
    sz3 = OxmlElement("w:sz")
    sz3.set(qn("w:val"), str(int(size * 2)))
    szCs3 = OxmlElement("w:szCs")
    szCs3.set(qn("w:val"), str(int(size * 2)))
    rPr3.append(sz3)
    rPr3.append(szCs3)
    r.append(rPr3)
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    r.append(fc2)
    para._p.append(r)


def _remove_table_borders(tbl):
    """Remove all borders from a table."""
    tbl_xml = tbl._tbl
    # Get or create tblPr
    tblPr = tbl_xml.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_xml.insert(0, tblPr)
    # Remove existing tblBorders if present
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        borders.append(b)
    tblPr.append(borders)


def _remove_para(para):
    """Delete a paragraph element from its parent."""
    p = para._element
    p.getparent().remove(p)


# ---------------------------------------------------------------------------
# Helpers: Formatting utilities
# ---------------------------------------------------------------------------

def _format_date(date_str):
    """Convert 'YYYY-MM-DD' to 'Month D, YYYY'."""
    if not date_str:
        return ""
    try:
        dt = datetime.strptime(date_str, "%Y-%m-%d")
        return f"{dt.strftime('%B')} {dt.day}, {dt.year}"
    except ValueError:
        return date_str


def _parse_fee(val):
    """Parse a fee value (int, string like '$12,960', or 0) to int."""
    if val is None:
        return 0
    if isinstance(val, (int, float)):
        return int(val)
    cleaned = re.sub(r"[^\d]", "", str(val))
    return int(cleaned) if cleaned else 0


# ---------------------------------------------------------------------------
# Fix 4: Word list numbering helpers
# ---------------------------------------------------------------------------

def _get_or_create_numbering_xml(doc):
    """Return the w:numbering XML element, creating the numbering part if needed."""
    try:
        np = doc.part.numbering_part
        if np is not None:
            return np._element
    except Exception:
        pass

    # Build a minimal numbering XML blob
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    blob = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        b"</w:numbering>"
    )
    content_type = (
        "application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.numbering+xml"
    )
    partname = PackURI("/word/numbering.xml")

    from docx.parts.numbering import NumberingPart as _NP
    numbering_part = _NP(partname, content_type, blob, doc.part.package)
    doc.part.relate_to(numbering_part, RT.NUMBERING)
    return numbering_part._element


def _add_abstract_num_def(nxml, abstract_id, num_fmt, lvl_text,
                          left_twips, hanging_twips=0):
    """Append a single-level abstractNum to nxml."""
    abs_num = OxmlElement("w:abstractNum")
    abs_num.set(qn("w:abstractNumId"), str(abstract_id))

    mlt = OxmlElement("w:multiLevelType")
    mlt.set(qn("w:val"), "hybridMultilevel")
    abs_num.append(mlt)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    lvl.append(fmt)

    txt = OxmlElement("w:lvlText")
    txt.set(qn("w:val"), lvl_text)
    lvl.append(txt)

    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)

    # Indentation in the abstract definition
    pPr_el = OxmlElement("w:pPr")
    ind_el = OxmlElement("w:ind")
    ind_el.set(qn("w:left"), str(left_twips))
    ind_el.set(qn("w:hanging"), str(hanging_twips))
    pPr_el.append(ind_el)
    lvl.append(pPr_el)

    rPr_el = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rPr_el.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")   # 9 pt
    rPr_el.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "18")
    rPr_el.append(szCs)
    lvl.append(rPr_el)

    abs_num.append(lvl)
    nxml.append(abs_num)


def _add_concrete_num(nxml, num_id, abstract_id):
    """Append a concrete <w:num> with startOverride=1 so each instance resets."""
    num_el = OxmlElement("w:num")
    num_el.set(qn("w:numId"), str(num_id))

    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num_el.append(abs_ref)

    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_el = OxmlElement("w:startOverride")
    start_el.set(qn("w:val"), "1")
    override.append(start_el)
    num_el.append(override)

    nxml.append(num_el)


def _setup_numbering(doc):
    """
    Create numbering abstract definitions and return a state dict used to
    allocate concrete numIds during document generation.
    Abstract 0 → uppercase letters (A, B, C …) at 0.5" indent
    Abstract 1 → decimal numbers (1. 2. 3. …) at 1.0" indent
    """
    nxml = _get_or_create_numbering_xml(doc)

    # Letters: "A)" label at 0.5", text at 1.0"  → left=1440, hanging=720
    # Numbers: "1."  label at 1.0", text at 1.5" → left=2160, hanging=720
    _add_abstract_num_def(nxml, abstract_id=0,
                          num_fmt="upperLetter", lvl_text="%1)",
                          left_twips=1440, hanging_twips=720)
    _add_abstract_num_def(nxml, abstract_id=1,
                          num_fmt="decimal", lvl_text="%1)",
                          left_twips=2160, hanging_twips=720)

    return {"nxml": nxml, "next_id": 1,
            "letter_abstract": 0, "number_abstract": 1}


def _alloc_letter_num(state):
    """Allocate a new concrete numId for a lettered list (resets to A)."""
    num_id = state["next_id"]
    state["next_id"] += 1
    _add_concrete_num(state["nxml"], num_id, state["letter_abstract"])
    return num_id


def _alloc_number_num(state):
    """Allocate a new concrete numId for a numbered list (resets to 1)."""
    num_id = state["next_id"]
    state["next_id"] += 1
    _add_concrete_num(state["nxml"], num_id, state["number_abstract"])
    return num_id


def _add_list_para(doc, html_content, num_id, left_inches=1.0, hanging_inches=0.5):
    """
    Add a list paragraph that uses Word's auto-numbering.
    label appears at (left_inches - hanging_inches), text starts at left_inches.
    Example: left=1.0", hanging=0.5" → label at 0.5", text at 1.0".
    """
    p = doc.add_paragraph()
    _para_spacing(p, 0, 0)

    pPr = p._p.get_or_add_pPr()

    # numPr must come first inside pPr per OOXML schema
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), "0")
    numPr.append(ilvl_el)
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(numId_el)
    pPr.insert(0, numPr)

    # Paragraph-level indent (overrides abstract num)
    ind_el = OxmlElement("w:ind")
    ind_el.set(qn("w:left"), str(int(left_inches * 1440)))
    ind_el.set(qn("w:hanging"), str(int(hanging_inches * 1440)))
    pPr.append(ind_el)

    # Process HTML body content (no label prefix run)
    soup = BeautifulSoup(html_content, "html.parser")
    _process_nodes(p, soup.children)

    # Handle nested <ul> bullet sub-lists
    for ul in soup.find_all("ul"):
        for li in ul.find_all("li"):
            bp = doc.add_paragraph()
            _para_spacing(bp, 0, 0)
            bpf = bp.paragraph_format
            bpf.left_indent = Inches(left_inches + 0.25)
            bpf.first_line_indent = Inches(-0.2)
            br = bp.add_run("\u2022\u2009")
            _fmt(br, 9)
            _process_nodes(bp, li.children)

    return p
